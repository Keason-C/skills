"""摄取:一个知识库根目录进去,知识目录 + 摄取报告出来。

深模块 —— 接口只有 `ingest_library` 一个函数,格式分派、扫描件识别、
品类推断、保密判定、生效日期抽取都藏在后面。

保密判定只看文档在不在保密文件夹下,没有第二判据(ADR-0005)。
没读进来的文件绝不静默丢弃,一律进摄取报告。
"""

from __future__ import annotations

import datetime as dt
import re
from collections.abc import Callable
from pathlib import Path

from .model import Document, IngestReport, Library, SkipReason, SkippedFile

DEFAULT_RESTRICTED_DIR = "保密"
UNCATEGORISED = "未分类"

SUPPORTED_SUFFIXES = {".docx", ".pdf", ".xlsx", ".xlsm", ".md", ".txt"}

# 只认"贴了标签的"日期。没贴标签的裸日期不当生效日期 —— 正文里的日期太多,
# 猜错比不猜更糟(排序会退回按摄取时间,见 answering)。
_DATE_LABELS = "生效日期|生效时间|生效|版本日期|签署日期|签订日期|日期"
_DATE_BODY = r"(\d{4})\s*[-/.年]\s*(\d{1,2})\s*[-/.月]\s*(\d{1,2})\s*日?"
_LABELLED_DATE = re.compile(rf"(?:{_DATE_LABELS})\s*[::]?\s*{_DATE_BODY}")


def ingest_library(
    root: Path | str,
    *,
    restricted_dir: str = DEFAULT_RESTRICTED_DIR,
    summariser: Callable[[Document], str] | None = None,
) -> Library:
    """走一遍知识库,返回知识目录与摄取报告。

    `summariser` 是可选的 `Callable[[Document], str]`,用来给每份文档生成一句话
    摘要(见 summarise.py)。不传就没有摘要,摄取照常完成。
    """
    root = Path(root)
    documents: list[Document] = []
    skipped: list[SkippedFile] = []

    for path in sorted(p for p in root.rglob("*") if p.is_file()):
        rel = path.relative_to(root).as_posix()
        suffix = path.suffix.lower()

        if suffix not in SUPPORTED_SUFFIXES:
            skipped.append(
                SkippedFile(rel, SkipReason.UNSUPPORTED_FORMAT, f"不支持的文件格式({suffix or '无扩展名'})")
            )
            continue

        try:
            text = _extract_text(path, suffix)
        except Exception as exc:  # 坏文件、加密文件、损坏的压缩包……
            skipped.append(SkippedFile(rel, SkipReason.UNREADABLE, f"读取失败:{type(exc).__name__}"))
            continue

        if not text.strip():
            if suffix == ".pdf":
                skipped.append(
                    SkippedFile(rel, SkipReason.SCANNED_PDF, "PDF 里没有文字层,疑似扫描件(本版不做文字识别)")
                )
            else:
                skipped.append(SkippedFile(rel, SkipReason.EMPTY, "文件里没有可读文字"))
            continue

        restricted, category = _classify(path.relative_to(root), restricted_dir)
        documents.append(
            Document(
                document_id=rel,
                title=path.stem,
                category=category,
                restricted=restricted,
                text=text,
                source_path=str(path),
                effective_date=_effective_date(text),
            )
        )

    if summariser is not None:
        documents = [_with_summary(doc, summariser) for doc in documents]

    return Library(
        documents=tuple(documents),
        report=IngestReport(
            ingested=tuple(d.document_id for d in documents),
            skipped=tuple(skipped),
        ),
    )


def _with_summary(doc: Document, summariser: Callable[[Document], str]) -> Document:
    """摘要生成失败不该让整次摄取失败 —— 没有摘要的文档照样进目录。"""
    from dataclasses import replace

    try:
        summary = (summariser(doc) or "").strip()
    except Exception:
        summary = ""
    return replace(doc, summary=summary)


def _classify(rel: Path, restricted_dir: str) -> tuple[bool, str]:
    parts = rel.parts
    if parts and parts[0] == restricted_dir:
        rest = parts[1:]
        category = rest[0] if len(rest) > 1 else UNCATEGORISED
        return True, category
    category = parts[0] if len(parts) > 1 else UNCATEGORISED
    return False, category


def _effective_date(text: str) -> dt.date | None:
    match = _LABELLED_DATE.search(text)
    if not match:
        return None
    year, month, day = (int(g) for g in match.groups())
    try:
        return dt.date(year, month, day)
    except ValueError:
        return None


def _extract_text(path: Path, suffix: str) -> str:
    if suffix in {".md", ".txt"}:
        return path.read_text(encoding="utf-8", errors="replace")
    if suffix == ".docx":
        return _extract_docx(path)
    if suffix in {".xlsx", ".xlsm"}:
        return _extract_xlsx(path)
    if suffix == ".pdf":
        return _extract_pdf(path)
    raise ValueError(f"unsupported suffix {suffix}")


def _extract_docx(path: Path) -> str:
    from docx import Document as DocxDocument

    doc = DocxDocument(str(path))
    lines = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            lines.append(" | ".join(cell.text for cell in row.cells))
    return "\n".join(lines)


def _extract_xlsx(path: Path) -> str:
    from openpyxl import load_workbook

    wb = load_workbook(str(path), data_only=True, read_only=True)
    lines: list[str] = []
    for ws in wb.worksheets:
        lines.append(f"[工作表] {ws.title}")
        for row in ws.iter_rows(values_only=True):
            cells = ["" if v is None else str(v) for v in row]
            if any(c.strip() for c in cells):
                lines.append(" | ".join(cells))
    wb.close()
    return "\n".join(lines)


def _extract_pdf(path: Path) -> str:
    from pypdf import PdfReader

    reader = PdfReader(str(path))
    return "\n".join((page.extract_text() or "") for page in reader.pages)
