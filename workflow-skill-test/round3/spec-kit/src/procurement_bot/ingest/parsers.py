"""五种格式的解析器。

契约:每个解析器返回 ``list[RawSection]``;**任何失败都抛异常**,由 loader 统一
捕获并转成 ``RejectedFile``。解析器自己绝不吞异常——静默丢失文档是宪法原则 III
明令禁止的。
"""

from __future__ import annotations

import re
from dataclasses import dataclass
from pathlib import Path

from ..errors import ParseError, UnsupportedFormatError


@dataclass(frozen=True)
class RawSection:
    heading_path: tuple[str, ...]
    locator: str
    text: str


SUPPORTED = {".md", ".txt", ".docx", ".pdf", ".xlsx"}


# ---------------------------------------------------------------------------
# 文本类
# ---------------------------------------------------------------------------

_MD_HEADING = re.compile(r"^(#{1,6})\s+(.*)$")


def parse_md(path: Path) -> list[RawSection]:
    text = path.read_text(encoding="utf-8", errors="replace")
    sections: list[RawSection] = []
    stack: list[tuple[int, str]] = []  # (级别, 标题) —— 不能假设"级别 == 下标+1"
    buf: list[str] = []

    def flush() -> None:
        body = "\n".join(buf).strip()
        if body:
            hp = tuple(title for _, title in stack)
            sections.append(
                RawSection(
                    heading_path=hp,
                    locator="§ " + " / ".join(hp) if hp else "正文",
                    text=body,
                )
            )
        buf.clear()

    for line in text.splitlines():
        m = _MD_HEADING.match(line)
        if m:
            flush()
            level = len(m.group(1))
            while stack and stack[-1][0] >= level:
                stack.pop()
            stack.append((level, m.group(2).strip()))
        else:
            buf.append(line)
    flush()
    return sections


def parse_txt(path: Path) -> list[RawSection]:
    text = path.read_text(encoding="utf-8", errors="replace").strip()
    if not text:
        return []
    return [RawSection(heading_path=(), locator="正文", text=text)]


# ---------------------------------------------------------------------------
# 二进制类
# ---------------------------------------------------------------------------


def parse_docx(path: Path) -> list[RawSection]:
    try:
        import docx  # type: ignore[import-untyped]
    except ImportError as exc:  # pragma: no cover - 依赖缺失时的兜底
        raise ParseError(f"缺少 python-docx: {exc}") from exc
    try:
        document = docx.Document(str(path))
    except Exception as exc:  # noqa: BLE001 - 任何打不开的情况都算解析失败
        raise ParseError(str(exc)) from exc

    sections: list[RawSection] = []
    stack: list[tuple[int, str]] = []
    buf: list[str] = []

    def flush() -> None:
        body = "\n".join(buf).strip()
        if body:
            hp = tuple(title for _, title in stack)
            sections.append(
                RawSection(
                    heading_path=hp,
                    locator="§ " + " / ".join(hp) if hp else "正文",
                    text=body,
                )
            )
        buf.clear()

    for para in document.paragraphs:
        content = para.text.strip()
        if not content:
            continue
        style = (para.style.name or "") if para.style is not None else ""
        m = re.match(r"Heading (\d+)", style)
        if m:
            flush()
            level = int(m.group(1))
            while stack and stack[-1][0] >= level:
                stack.pop()
            stack.append((level, content))
        else:
            buf.append(content)
    flush()

    for i, table in enumerate(document.tables, start=1):
        rows = [
            " | ".join(cell.text.strip() for cell in row.cells) for row in table.rows
        ]
        body = "\n".join(r for r in rows if r.strip(" |"))
        if body:
            sections.append(
                RawSection(heading_path=("表格",), locator=f"表格 {i}", text=body)
            )

    return sections


def parse_pdf(path: Path) -> list[RawSection]:
    import logging

    try:
        import pypdf
    except ImportError as exc:  # pragma: no cover
        raise ParseError(f"缺少 pypdf: {exc}") from exc

    # pypdf 会把"文件坏了"直接 print 到 stderr。我们自己会把它转成
    # RejectedFile 报给用户,不需要库再吵一遍。
    pypdf_logger = logging.getLogger("pypdf")
    previous = pypdf_logger.level
    pypdf_logger.setLevel(logging.CRITICAL)
    try:
        reader = pypdf.PdfReader(str(path))
        pages = list(reader.pages)
    except Exception as exc:  # noqa: BLE001
        raise ParseError(str(exc)) from exc
    finally:
        pypdf_logger.setLevel(previous)

    sections: list[RawSection] = []
    for i, page in enumerate(pages, start=1):
        try:
            text = (page.extract_text() or "").strip()
        except Exception as exc:  # noqa: BLE001
            raise ParseError(f"第 {i} 页无法抽取文本: {exc}") from exc
        if text:
            sections.append(
                RawSection(heading_path=(), locator=f"第 {i} 页", text=text)
            )

    if not sections:
        # 扫描件 / 纯图片 PDF:不静默跳过,明确报告(宪法原则 III)
        raise ParseError("PDF 中没有可抽取的文字(可能是扫描件或纯图片)")
    return sections


def parse_xlsx(path: Path) -> list[RawSection]:
    try:
        import openpyxl
    except ImportError as exc:  # pragma: no cover
        raise ParseError(f"缺少 openpyxl: {exc}") from exc
    try:
        wb = openpyxl.load_workbook(str(path), read_only=True, data_only=True)
    except Exception as exc:  # noqa: BLE001
        raise ParseError(str(exc)) from exc

    sections: list[RawSection] = []
    try:
        for ws in wb.worksheets:
            rows = list(ws.iter_rows(values_only=True))
            if not rows:
                continue
            header = [str(c).strip() if c is not None else "" for c in rows[0]]
            # 每行都会被渲染成"表头: 值"的自足形式,所以不再单独输出一遍表头行——
            # 否则检索最容易命中的会是那行没有信息量的表头。
            lines: list[str] = []
            for row in rows[1:]:
                cells = [str(c).strip() if c is not None else "" for c in row]
                if not any(cells):
                    continue
                pairs = [
                    f"{header[i]}: {cells[i]}"
                    for i in range(min(len(header), len(cells)))
                    if cells[i] and (header[i] if i < len(header) else "")
                ]
                # 用 " | " 而不是分号:分号是句子分隔符,会把一行数据切碎
                lines.append(" | ".join(pairs) if pairs else " | ".join(cells))
            if not lines:  # 只有表头、没有数据行时,至少把表头留下
                lines = [" | ".join(h for h in header if h)]
            body = "\n".join(line for line in lines if line.strip())
            if body:
                sections.append(
                    RawSection(
                        heading_path=(ws.title,),
                        locator=f"工作表:{ws.title}",
                        text=body,
                    )
                )
    finally:
        wb.close()

    return sections


# ---------------------------------------------------------------------------
# 调度
# ---------------------------------------------------------------------------

_PARSERS = {
    ".md": parse_md,
    ".markdown": parse_md,
    ".txt": parse_txt,
    ".docx": parse_docx,
    ".pdf": parse_pdf,
    ".xlsx": parse_xlsx,
}


def parse(path: Path) -> list[RawSection]:
    ext = path.suffix.lower()
    parser = _PARSERS.get(ext)
    if parser is None:
        raise UnsupportedFormatError(f"不支持的格式({ext or '无扩展名'})")
    return parser(path)
