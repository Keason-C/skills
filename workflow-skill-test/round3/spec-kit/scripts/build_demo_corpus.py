#!/usr/bin/env python3
"""生成演示语料 —— 一套**假的**采购文档,中英混排。

需求方原话:"样例文档现在拿不出脱敏版,你先自己造一套像样的假采购文档
(品类、供应商、规格、流程,外加保密的报价/成本/谈判纪要,中英混着),
保密验收就用假的验。"

里面的公司名、人名、价格全是编的。

PDF 用一个自写的最小生成器(约 40 行,手工组装 xref 表),
这样就不用为了造测试数据多引入一个 reportlab 依赖(research.md D-03)。
副作用:PDF 只能是英文(base14 字体不含中文)——正好满足"中英混着"。
"""

from __future__ import annotations

import shutil
from pathlib import Path

ROOT = Path(__file__).resolve().parent.parent
CORPUS = ROOT / "demo" / "corpus"
SECRET = CORPUS / "保密"


# ---------------------------------------------------------------------------
# 最小 PDF 生成器(无第三方依赖)
# ---------------------------------------------------------------------------


def make_pdf(lines: list[str]) -> bytes:
    def esc(s: str) -> str:
        return s.replace("\\", "\\\\").replace("(", "\\(").replace(")", "\\)")

    content = (
        "BT /F1 11 Tf 60 780 Td 16 TL\n"
        + "".join(f"({esc(line)}) Tj T*\n" for line in lines)
        + "ET"
    )
    cb = content.encode("latin-1", errors="replace")
    objs = [
        b"<</Type/Catalog/Pages 2 0 R>>",
        b"<</Type/Pages/Kids[3 0 R]/Count 1>>",
        b"<</Type/Page/Parent 2 0 R/MediaBox[0 0 595 842]"
        b"/Resources<</Font<</F1 5 0 R>>>>/Contents 4 0 R>>",
        b"<</Length %d>>stream\n" % len(cb) + cb + b"\nendstream",
        b"<</Type/Font/Subtype/Type1/BaseFont/Helvetica>>",
    ]
    out = bytearray(b"%PDF-1.4\n")
    offsets: list[int] = []
    for i, body in enumerate(objs, start=1):
        offsets.append(len(out))
        out += b"%d 0 obj\n" % i + body + b"\nendobj\n"
    xref = len(out)
    out += b"xref\n0 %d\n0000000000 65535 f \n" % (len(objs) + 1)
    for off in offsets:
        out += b"%010d 00000 n \n" % off
    out += b"trailer\n<</Size %d/Root 1 0 R>>\nstartxref\n%d\n%%%%EOF\n" % (
        len(objs) + 1,
        xref,
    )
    return bytes(out)


# ---------------------------------------------------------------------------


def write_text(path: Path, text: str) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(text, encoding="utf-8")
    print(f"  + {path.relative_to(ROOT)}")


def write_bytes(path: Path, data: bytes) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_bytes(data)
    print(f"  + {path.relative_to(ROOT)}")


def build_docx(path: Path, title: str, blocks: list[tuple[str, str]]) -> None:
    import docx

    d = docx.Document()
    d.add_heading(title, level=1)
    for style, text in blocks:
        if style == "h2":
            d.add_heading(text, level=2)
        else:
            d.add_paragraph(text)
    path.parent.mkdir(parents=True, exist_ok=True)
    d.save(str(path))
    print(f"  + {path.relative_to(ROOT)}")


def build_xlsx(path: Path, sheets: dict[str, list[list[str]]]) -> None:
    import openpyxl

    wb = openpyxl.Workbook()
    wb.remove(wb.active)
    for name, rows in sheets.items():
        ws = wb.create_sheet(title=name)
        for row in rows:
            ws.append(row)
    path.parent.mkdir(parents=True, exist_ok=True)
    wb.save(str(path))
    print(f"  + {path.relative_to(ROOT)}")


def main() -> None:
    if CORPUS.exists():
        shutil.rmtree(CORPUS)
    print("生成演示语料(全部为虚构内容):")

    # ---- 普通文档 -------------------------------------------------------
    write_text(
        CORPUS / "品类管理手册-电解铜.md",
        """# 电解铜品类管理手册

## 品类归属

本品类负责人:张伟(工号 P1002),分机 8021。品类英文代号 COPPER-CATHODE。
备份负责人:赵强(工号 P1005)。

## 规格要求

电解铜执行 Cu-CATH-1 标准(GB/T 467-2010),含铜量不低于 99.95%,
表面不得有明显氧化斑。每批次须附第三方检测报告。

## 常见问题

问:电解铜的采购周期多久?答:常规订单 30 天,加急 15 天(需品类经理审批)。
""",
    )

    write_text(
        CORPUS / "供应商准入流程.txt",
        """供应商准入流程

第一步 资质初审:营业执照、生产许可、ISO9001 证书。
第二步 现场审核:由品类经理带队,质量部参加,出具审核报告。
第三步 样品检测:送第三方实验室,合格后方可进入下一步。
第四步 列入合格供应商清单:由采购负责人签批。

准入周期通常为 6 至 8 周。
""",
    )

    # 两份互相矛盾、且都没有日期 —— 触发 FR-005b(并列摆出)
    write_text(
        CORPUS / "规格要求-紧固件.md",
        """# 紧固件规格要求

## 表面处理

紧固件表面处理采用达克罗(Dacromet),盐雾试验不低于 720 小时。

## 材质

主体材质为 SCM435 合金钢。
""",
    )
    write_text(
        CORPUS / "规格要求-紧固件-补充.md",
        """# 紧固件规格要求补充

## 表面处理

紧固件表面处理一律采用镀锌钝化,盐雾试验不低于 480 小时。

## 材质

主体材质为 SCM435 合金钢。
""",
    )

    # ---- Word:两个版本,都有生效日期 —— 触发 FR-005a(取新 + 注明旧版) ----
    build_docx(
        CORPUS / "采购流程说明.docx",
        "采购流程说明",
        [
            ("p", "生效日期:2024-03-01"),
            ("h2", "审批环节"),
            ("p", "审批流程共三级:主管、经理、总监。10 万元以下由主管审批即可。"),
            ("h2", "请购"),
            ("p", "请购单在系统中提交,须注明物料编码、数量、需求日期。"),
        ],
    )
    build_docx(
        CORPUS / "采购流程说明-2019旧版.docx",
        "采购流程说明",
        [
            ("p", "生效日期:2019-05-01"),
            ("h2", "审批环节"),
            ("p", "审批流程共五级:主管、经理、总监、财务、总经理。"),
        ],
    )

    # ---- Excel:多工作表 -------------------------------------------------
    build_xlsx(
        CORPUS / "供应商清单-电解铜.xlsx",
        {
            "合格供应商": [
                ["供应商名称", "供应商编码", "所在地", "年供货量(吨)", "准入日期"],
                ["江铜国贸(虚构)", "SUP-0101", "江西鹰潭", "12000", "2021-04-12"],
                ["铜冠物流(虚构)", "SUP-0102", "安徽铜陵", "8000", "2022-09-01"],
                ["紫金供应链(虚构)", "SUP-0103", "福建上杭", "5000", "2023-02-20"],
            ],
            "备选供应商": [
                ["供应商名称", "供应商编码", "状态"],
                ["中原有色(虚构)", "SUP-0201", "现场审核通过,待样品检测"],
            ],
        },
    )

    # ---- PDF:英文(中英混排的"英"这一半) ------------------------------
    write_bytes(
        CORPUS / "Category_Ownership_Matrix.pdf",
        make_pdf(
            [
                "Category Ownership Matrix (FY2026)",
                "",
                "Category            Owner            Employee ID   Backup",
                "Copper Cathode      Zhang Wei        P1002         Zhao Qiang",
                "Fasteners           Wang Fang        P1003         -",
                "Aluminium Ingot     Zhao Qiang       P1005         -",
                "",
                "Escalation: any unassigned category goes to Chen Guohua (P9000).",
            ]
        ),
    )

    # ---- 读不了的文件:必须出现在"未纳入清单"里 --------------------------
    write_bytes(CORPUS / "扫描件-旧合同.pdf", b"%PDF-1.4\n truncated and corrupted \n")
    write_bytes(CORPUS / "说明.bin", b"\x00\x01\x02 binary blob \xff\xfe")

    # ---- 保密文档 -------------------------------------------------------
    write_text(
        SECRET / "电解铜" / "成本构成说明.md",
        """# 电解铜成本构成(保密)

## 采购单价

2025 年协议采购单价为 68500 元/吨,其中加工费 1200 元/吨,物流费 380 元/吨。

## 成本结构

LME 基准 + 升水。当前升水 95 美元/吨,较去年下降 12 美元/吨。
""",
    )
    write_text(
        SECRET / "电解铜" / "谈判纪要-2025Q1.md",
        """# 电解铜谈判纪要 2025Q1(保密)

## 谈判策略

以年度总量 12000 吨为筹码,要求升水下调至 88 美元/吨。
底线:升水不高于 98 美元/吨,否则转备选供应商。
""",
    )
    write_text(
        SECRET / "紧固件" / "谈判策略要点.md",
        """# 紧固件谈判策略要点(保密)

## 底线

年降目标 5%,可让步至 3%。低于 3% 需报采购负责人审批。
供应商 SUP-0301 的实际成本约为报价的 72%。
""",
    )
    # 未归品类:只有采购负责人能看(FR-008a)
    write_text(
        SECRET / "集团年度成本目标.md",
        """# 集团年度采购降本目标(保密)

全集团 2026 年度采购降本目标为 7.2%,其中原材料板块 8.5%。
""",
    )

    # ---- 人员名单 -------------------------------------------------------
    roster = ROOT / "demo" / "roster.csv"
    roster.parent.mkdir(parents=True, exist_ok=True)
    roster.write_text(
        "工号,姓名,角色,负责品类\n"
        "P1001,李明,采购员,\n"
        "P1002,张伟,品类经理,电解铜\n"
        "P1003,王芳,品类经理,紧固件\n"
        "P1005,赵强,品类经理,电解铜;铝锭\n"
        "P9000,陈国华,采购负责人,*\n",
        encoding="utf-8",
    )
    print(f"  + {roster.relative_to(ROOT)}")
    print("\n完成。接下来:")
    print("  pbot ingest --corpus demo/corpus --roster demo/roster.csv --state .pbot")


if __name__ == "__main__":
    main()
