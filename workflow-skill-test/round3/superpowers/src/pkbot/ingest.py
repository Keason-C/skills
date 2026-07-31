"""把各种格式的文件变成带出处标记的文本小节。读不了就明说,绝不猜内容。

需求方原话:"大文件读不动宁可说读不动"、"读不了的文件明确列出来别装懂"。
"""

from __future__ import annotations

import csv
from pathlib import Path

from .models import Section

MAX_ROWS_PER_SECTION = 200


class UnreadableError(Exception):
    """这份文件读不了,消息体是给非技术同事看的原因。"""


def _sections_from_markdown(text: str) -> tuple[Section, ...]:
    sections: list[Section] = []
    title = "全文"
    buf: list[str] = []

    def flush() -> None:
        if buf and any(line.strip() for line in buf):
            sections.append(Section(title, "\n".join(buf).strip()))

    for line in text.splitlines():
        if line.startswith("#"):
            flush()
            title = line.lstrip("#").strip() or "全文"
            buf = []
        else:
            buf.append(line)
    flush()
    return tuple(sections)


def _parse_text(path: Path) -> tuple[Section, ...]:
    text = path.read_text(encoding="utf-8", errors="replace")
    if not text.strip():
        return ()
    if path.suffix.lower() == ".md":
        sections = _sections_from_markdown(text)
        if sections:
            return sections
    return (Section("全文", text.strip()),)


def _parse_docx(path: Path) -> tuple[Section, ...]:
    from docx import Document as DocxDocument

    doc = DocxDocument(str(path))
    sections: list[Section] = []
    title = "全文"
    buf: list[str] = []

    def flush() -> None:
        if buf and any(line.strip() for line in buf):
            sections.append(Section(title, "\n".join(buf).strip()))

    for para in doc.paragraphs:
        style_name = para.style.name if para.style is not None else ""
        is_heading = style_name.startswith("Heading") or style_name.startswith("标题")
        if is_heading and para.text.strip():
            flush()
            title = para.text.strip()
            buf = []
        elif para.text.strip():
            buf.append(para.text)
    flush()

    for idx, table in enumerate(doc.tables, start=1):
        rows = ["\t".join(cell.text.strip() for cell in row.cells) for row in table.rows]
        body = "\n".join(row for row in rows if row.strip())
        if body:
            sections.append(Section(f"表格 {idx}", body))
    return tuple(sections)


def _rows_to_sections(label: str, rows: list[list[str]]) -> list[Section]:
    """把表格行按固定块大小切片,每片带明确的行号范围作为出处。"""
    out: list[Section] = []
    for start in range(0, len(rows), MAX_ROWS_PER_SECTION):
        chunk = rows[start : start + MAX_ROWS_PER_SECTION]
        body = "\n".join(
            "\t".join(cell for cell in row) for row in chunk if any(c.strip() for c in row)
        )
        if body:
            out.append(Section(f"{label} 第 {start + 1}-{start + len(chunk)} 行", body))
    return out


def _parse_xlsx(path: Path) -> tuple[Section, ...]:
    from openpyxl import load_workbook

    wb = load_workbook(str(path), read_only=True, data_only=True)
    try:
        sections: list[Section] = []
        for ws in wb.worksheets:
            rows = [
                ["" if cell is None else str(cell) for cell in row]
                for row in ws.iter_rows(values_only=True)
            ]
            sections.extend(_rows_to_sections(f"工作表 {ws.title}", rows))
    finally:
        wb.close()
    return tuple(sections)


def _parse_csv(path: Path) -> tuple[Section, ...]:
    with path.open(encoding="utf-8-sig", newline="") as fh:
        rows = [list(row) for row in csv.reader(fh)]
    return tuple(_rows_to_sections("表格", rows))


def _parse_pdf(path: Path) -> tuple[Section, ...]:
    from pypdf import PdfReader
    from pypdf.errors import PdfReadError

    try:
        reader = PdfReader(str(path))
        if reader.is_encrypted:
            raise UnreadableError(f"{path.name}:PDF 已加密,打不开。")
        sections: list[Section] = []
        for idx, page in enumerate(reader.pages, start=1):
            text = (page.extract_text() or "").strip()
            if text:
                sections.append(Section(f"第 {idx} 页", text))
        return tuple(sections)
    except UnreadableError:
        raise
    except (PdfReadError, OSError, ValueError) as exc:
        raise UnreadableError(f"{path.name}:PDF 解析失败({exc})。") from exc


_PARSERS = {
    ".md": _parse_text,
    ".txt": _parse_text,
    ".docx": _parse_docx,
    ".xlsx": _parse_xlsx,
    ".csv": _parse_csv,
    ".pdf": _parse_pdf,
}

SUPPORTED_SUFFIXES = set(_PARSERS)


def parse_sections(path: Path) -> tuple[Section, ...]:
    parser = _PARSERS.get(path.suffix.lower())
    if parser is None:
        raise UnreadableError(
            f"{path.name}:不支持的格式 {path.suffix or '(无扩展名)'},"
            f"目前只认 {'、'.join(sorted(SUPPORTED_SUFFIXES))}。"
        )
    try:
        sections = parser(path)
    except UnreadableError:
        raise
    except Exception as exc:  # 解析库异常五花八门,统一转成人话
        raise UnreadableError(
            f"{path.name}:读取失败({type(exc).__name__}: {exc})。"
        ) from exc
    if not sections:
        raise UnreadableError(f"{path.name}:没有可提取的文字(可能是扫描件或空文件)。")
    return sections
