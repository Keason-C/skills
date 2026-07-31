"""生成真实的 .docx / .xlsx / .pdf fixture 文件。

刻意不用字符串假装解析:解析器的价值就在于能真的打开这些二进制格式。
"""

import pytest
from docx import Document as DocxDocument
from openpyxl import Workbook
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas


@pytest.fixture
def docx_file(tmp_path):
    p = tmp_path / "谈判纪要.docx"
    doc = DocxDocument()
    doc.add_heading("一、背景", level=1)
    doc.add_paragraph("2025 年度紧固件年框谈判。")
    doc.add_heading("二、结论", level=1)
    doc.add_paragraph("单价下降 3%。")
    doc.save(p)
    return p


@pytest.fixture
def xlsx_file(tmp_path):
    p = tmp_path / "对照表.xlsx"
    wb = Workbook()
    ws = wb.active
    ws.title = "品类负责人"
    ws.append(["品类", "负责人", "工号"])
    ws.append(["紧固件", "王五", "G0007"])
    wb.save(p)
    return p


@pytest.fixture
def pdf_file(tmp_path):
    # PDF fixture 内容用 ASCII,避免依赖 CJK 字体;
    # 解析器是格式级的,内容语言与它是否正确无关。
    p = tmp_path / "tender.pdf"
    c = canvas.Canvas(str(p), pagesize=A4)
    c.drawString(72, 720, "SECTION 3 TECHNICAL SPEC")
    c.drawString(72, 700, "Bolt M8x30 GB/T 5783")
    c.showPage()
    c.drawString(72, 720, "SECTION 4 DELIVERY TERMS")
    c.showPage()
    c.save()
    return p
